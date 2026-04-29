from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from weasyprint import HTML

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top', 'start', 'bottom', 'end'):
        tag = 'w:{}'.format(side)
        element = tcPr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcPr.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), '000000')

# Define layout based on the image provided
# Across paths
across = [
    (1, 8, 5, "1"),   # SCRUM
    (1, 13, 5, "4"),  # AGILE
    (3, 5, 6, "5"),   # DEVOPS
    (5, 15, 6, "6"),  # SPRINT
    (7, 4, 9, "8"),   # WATERFALL
    (10, 10, 8, "10"),# OUTSOURCE
    (12, 12, 9, "13"),# DASHBOARD
    (14, 5, 8, "15"), # FIREWALL
    (14, 15, 3, "16"),# CRM
    (17, 6, 6, "18")  # PORTAL
]

# Down paths
down = [
    (1, 9, 5, "2"),   # CLOUD
    (1, 18, 6, "3"),  # VENDOR
    (3, 5, 6, "5"),   # DATABASE (Starts same cell as DEVOPS)
    (5, 15, 6, "6"),  # SERVER (Starts same cell as SPRINT)
    (7, 10, 9, "7"),  # PROTOTYPE
    (7, 20, 7, "9"),  # NETWORK
    (10, 18, 8, "11"),# ECOMMERCE
    (12, 6, 3, "12"), # APP
    (13, 11, 6, "14"),# KANBAN
    (16, 10, 3, "17") # ERP
]

def generate_word():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    
    table = doc.add_table(rows=22, cols=25)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    active_cells = {} # (r, c) -> number

    for r, c, length, num in across:
        for i in range(length):
            pos = (r, c + i)
            if pos not in active_cells: active_cells[pos] = ""
            if i == 0: active_cells[pos] = num

    for r, c, length, num in down:
        for i in range(length):
            pos = (r + i, c)
            if pos not in active_cells: active_cells[pos] = ""
            if i == 0: active_cells[pos] = num

    for r in range(22):
        row = table.rows[r]
        row.height = Cm(0.7)
        for c in range(25):
            cell = table.cell(r, c)
            if (r, c) in active_cells:
                set_cell_border(cell)
                if active_cells[(r, c)]:
                    p = cell.paragraphs[0]
                    run = p.add_run(active_cells[(r, c)])
                    run.font.size = Pt(8)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save("empty_crossword_layout_v1.docx")

def generate_pdf():
    # Simple HTML table for PDF
    cells = {}
    for r, c, length, num in across:
        for i in range(length):
            pos = (r, c + i)
            if pos not in cells: cells[pos] = ""
            if i == 0: cells[pos] = num
    for r, c, length, num in down:
        for i in range(length):
            pos = (r + i, c)
            if pos not in cells: cells[pos] = ""
            if i == 0: cells[pos] = num

    html_content = """
    <html>
    <style>
        @page { size: A4; margin: 1cm; background-color: #ffffff; }
        table { border-collapse: collapse; margin: auto; }
        td { width: 25px; height: 25px; border: 1px solid #ddd; position: relative; font-family: sans-serif; }
        .active { border: 2px solid black; }
        .num { position: absolute; top: 1px; left: 2px; font-size: 8px; }
    </style>
    <body>
        <table>
    """
    for r in range(22):
        html_content += "<tr>"
        for c in range(25):
            if (r, c) in cells:
                num = cells[(r, c)]
                html_content += f'<td class="active"><span class="num">{num}</span></td>'
            else:
                html_content += '<td></td>'
        html_content += "</tr>"
    html_content += "</table></body></html>"
    
    with open("temp.html", "w") as f:
        f.write(html_content)
    HTML("temp.html").write_pdf("empty_crossword_layout_v1.pdf")

generate_word()
generate_pdf()